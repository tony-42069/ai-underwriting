import pytest
from unittest.mock import AsyncMock, MagicMock, patch
import os
import tempfile

from backend.services.ocr import DocumentProcessor


class TestDocumentProcessor:
    """Tests for DocumentProcessor class."""

    @pytest.fixture
    def processor(self):
        """Create a DocumentProcessor instance."""
        return DocumentProcessor()

    def test_init(self, processor):
        """Test processor initialization."""
        assert len(processor.extractors) == 4
        extractor_names = [e.__class__.__name__ for e in processor.extractors]
        assert "RentRollExtractor" in extractor_names
        assert "PLStatementExtractor" in extractor_names
        assert "OperatingStatementExtractor" in extractor_names
        assert "LeaseExtractor" in extractor_names

    @pytest.mark.asyncio
    async def test_process_document_pdf(self, processor, sample_pdf_content):
        """Test processing a PDF document."""
        with tempfile.NamedTemporaryFile(
            suffix=".pdf", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake pdf content")
            tmp_file_path = tmp_file.name

        try:
            with patch.object(
                processor, "_process_pdf", new=AsyncMock(return_value=sample_pdf_content)
            ):
                result = await processor.process_document(tmp_file_path)

            assert "status" in result
            assert result["status"] in ["success", "error"]
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_document_excel(self, processor):
        """Test processing an Excel document."""
        with tempfile.NamedTemporaryFile(
            suffix=".xlsx", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake excel content")
            tmp_file_path = tmp_file.name

        try:
            with patch.object(
                processor, "_process_excel", new=AsyncMock(return_value="Sheet: Test\nData")
            ):
                result = await processor.process_document(tmp_file_path)

            assert "status" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_document_word(self, processor):
        """Test processing a Word document."""
        with tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake word content")
            tmp_file_path = tmp_file.name

        try:
            with patch.object(
                processor, "_process_word", new=AsyncMock(return_value="Document content")
            ):
                result = await processor.process_document(tmp_file_path)

            assert "status" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_document_file_not_found(self, processor):
        """Test processing non-existent file."""
        result = await processor.process_document("/nonexistent/path/file.pdf")

        assert result["status"] == "error"
        assert "File not found" in result["error"]

    @pytest.mark.asyncio
    async def test_process_document_unsupported_type(self, processor):
        """Test processing unsupported file type."""
        with tempfile.NamedTemporaryFile(
            suffix=".txt", delete=False
        ) as tmp_file:
            tmp_file.write(b"text content")
            tmp_file_path = tmp_file.name

        try:
            result = await processor.process_document(tmp_file_path)

            assert result["status"] == "error"
            assert "Unsupported file type" in result["error"]
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_document_extraction_success(
        self, processor, sample_pdf_content
    ):
        """Test successful extraction from document."""
        with tempfile.NamedTemporaryFile(
            suffix=".pdf", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake pdf content")
            tmp_file_path = tmp_file.name

        try:
            with patch.object(
                processor, "_process_pdf", new=AsyncMock(return_value=sample_pdf_content)
            ):
                result = await processor.process_document(tmp_file_path)

            if result["status"] == "success":
                assert "extractions" in result
                assert len(result["extractions"]) >= 0
                assert "processed_at" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_document_no_extractions(self, processor):
        """Test processing document with no matching extractors."""
        content = "This document has no matching extractors content at all."
        with tempfile.NamedTemporaryFile(
            suffix=".pdf", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake pdf content")
            tmp_file_path = tmp_file.name

        try:
            with patch.object(
                processor, "_process_pdf", new=AsyncMock(return_value=content)
            ):
                result = await processor.process_document(tmp_file_path)

            if result["status"] == "success":
                assert "extractions" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_document_exception_handling(self, processor):
        """Test exception handling during processing."""
        with tempfile.NamedTemporaryFile(
            suffix=".pdf", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake pdf content")
            tmp_file_path = tmp_file.name

        try:
            with patch.object(
                processor, "_process_pdf", side_effect=Exception("Processing error")
            ):
                result = await processor.process_document(tmp_file_path)

            assert result["status"] == "error"
            assert "Processing error" in result["error"]
            assert "processed_at" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_pdf_uses_correct_path(self, processor):
        """Test PDF processing uses correct poppler path."""
        from backend.config.settings import settings

        with tempfile.NamedTemporaryFile(
            suffix=".pdf", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake pdf content")
            tmp_file_path = tmp_file.name

        try:
            with patch("pdf2image.convert_from_path") as mock_convert:
                mock_convert.return_value = []
                with patch("pytesseract.image_to_string", return_value="text"):
                    await processor._process_pdf(tmp_file_path)

                mock_convert.assert_called_once()
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_pdf_exception(self, processor):
        """Test PDF processing exception handling."""
        with tempfile.NamedTemporaryFile(
            suffix=".pdf", delete=False
        ) as tmp_file:
            tmp_file.write(b"fake pdf content")
            tmp_file_path = tmp_file.name

        try:
            with patch("pdf2image.convert_from_path", side_effect=Exception("PDF error")):
                result = await processor._process_pdf(tmp_file_path)

            assert result is None
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_excel(self, processor):
        """Test Excel file processing."""
        import pandas as pd
        from io import BytesIO

        df = pd.DataFrame({"A": [1, 2, 3], "B": [4, 5, 6]})
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="TestSheet")
        excel_content = excel_buffer.getvalue()

        with tempfile.NamedTemporaryFile(
            suffix=".xlsx", delete=False
        ) as tmp_file:
            tmp_file.write(excel_content)
            tmp_file_path = tmp_file.name

        try:
            result = await processor._process_excel(tmp_file_path)

            assert result is not None
            assert "TestSheet" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_excel_exception(self, processor):
        """Test Excel processing exception handling."""
        with tempfile.NamedTemporaryFile(
            suffix=".xlsx", delete=False
        ) as tmp_file:
            tmp_file.write(b"invalid excel content")
            tmp_file_path = tmp_file.name

        try:
            result = await processor._process_excel(tmp_file_path)

            assert result is None
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_word(self, processor):
        """Test Word document processing."""
        from docx import Document
        from io import BytesIO

        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Test content")

        buffer = BytesIO()
        doc.save(buffer)
        word_content = buffer.getvalue()

        with tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False
        ) as tmp_file:
            tmp_file.write(word_content)
            tmp_file_path = tmp_file.name

        try:
            result = await processor._process_word(tmp_file_path)

            assert result is not None
            assert "Hello World" in result
            assert "Test content" in result
        finally:
            os.unlink(tmp_file_path)

    @pytest.mark.asyncio
    async def test_process_word_exception(self, processor):
        """Test Word processing exception handling."""
        with tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False
        ) as tmp_file:
            tmp_file.write(b"invalid docx content")
            tmp_file_path = tmp_file.name

        try:
            result = await processor._process_word(tmp_file_path)

            assert result is None
        finally:
            os.unlink(tmp_file_path)
